from flask import Flask, render_template, request, redirect, session, flash, send_from_directory, url_for, make_response
import sqlite3
import os
import uuid
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import urllib.parse
import json

app = Flask(__name__)
app.secret_key = 'super_secret_key_pour_session'

# Cr√©er les dossiers si n√©cessaire
os.makedirs('uploads', exist_ok=True)
os.makedirs('rapports', exist_ok=True)
os.makedirs('static', exist_ok=True)

# Initialiser la base de donn√©es
def init_db():
    conn = sqlite3.connect('database.db')
    c = conn.cursor()
    
    # Table utilisateurs
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        email TEXT UNIQUE,
        password TEXT,
        full_name TEXT,
        role TEXT CHECK(role IN ('admin', 'technicien', 'medecin', 'agent_tireur'))
    )''')
    
    # Table patients
    c.execute('''CREATE TABLE IF NOT EXISTS patients (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nom TEXT NOT NULL,
        prenom TEXT NOT NULL,
        age INTEGER,
        diagnostic TEXT,
        examen_demande TEXT,
        id_patient TEXT UNIQUE NOT NULL,
        date_naissance TEXT,
        telephone TEXT
    )''')
    
    # Table examens
    c.execute('''CREATE TABLE IF NOT EXISTS exams (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_id INTEGER NOT NULL,
        technicien_id INTEGER,
        medecin_id INTEGER,
        agent_tireur_id INTEGER,
        status TEXT DEFAULT 'Upload√©',
        dicom_path TEXT,
        rapport_path TEXT,
        uploaded_at TEXT NOT NULL,
        visualized_at TEXT,
        rapport_submitted_at TEXT,
        rapport_printed_at TEXT,
        FOREIGN KEY (patient_id) REFERENCES patients (id),
        FOREIGN KEY (technicien_id) REFERENCES users (id),
        FOREIGN KEY (medecin_id) REFERENCES users (id),
        FOREIGN KEY (agent_tireur_id) REFERENCES users (id)
    )''')
    
    conn.commit()
    conn.close()

init_db()

# ‚úÖ Cr√©er les comptes par d√©faut
def create_default_users():
    conn = sqlite3.connect('database.db')
    c = conn.cursor()
    
    users = [
        ("admin@pdo.ci", "admin123", "Admin PDO", "admin"),
        ("drcamara@pdo.ci", "med123", "Dr Camara", "medecin"),
        ("drfrancistraore@pdo.ci", "med123", "Dr Francis Traore", "medecin"),
        ("agent@pdo.ci", "agent123", "Agent Tireur", "agent_tireur"),
        ("tuocyrille@pdo.ci", "tech123", "Technicien Cyrille", "technicien"),
        ("ayemou@pdo.ci", "tech123", "Technicien Ayemou", "technicien"),
        ("kone@pdo.ci", "tech123", "Technicien Kone", "technicien"),
    ]
    
    for email, password, full_name, role in users:
        hashed_pw = generate_password_hash(password)
        try:
            c.execute("INSERT INTO users (email, password, full_name, role) VALUES (?, ?, ?, ?)",
                      (email, hashed_pw, full_name, role))
        except sqlite3.IntegrityError:
            pass
    
    conn.commit()
    conn.close()

create_default_users()

@app.route('/')
def index():
    if 'user_id' not in session:
        return redirect('/login')
    role = session.get('role')
    if role == 'admin':
        return redirect('/clinique')
    elif role == 'technicien':
        return redirect('/technicien')
    elif role == 'medecin':
        return redirect('/medecin')
    elif role == 'agent_tireur':
        return redirect('/agent_tireur')
    return redirect('/login')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        c.execute("SELECT * FROM users WHERE email = ?", (email,))
        user = c.fetchone()
        conn.close()
        if user and check_password_hash(user[2], password):
            session['user_id'] = user[0]
            session['role'] = user[4]
            session['full_name'] = user[3]
            session['user_email'] = email
            flash('‚úÖ Connexion r√©ussie !', 'success')
            return redirect('/')
        else:
            flash('‚ùå Email ou mot de passe incorrect', 'error')
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form['email']
        password = generate_password_hash(request.form['password'])
        full_name = request.form['full_name']
        role = request.form['role']
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        try:
            c.execute("INSERT INTO users (email, password, full_name, role) VALUES (?, ?, ?, ?)",
                      (email, password, full_name, role))
            conn.commit()
            flash('‚úÖ Inscription r√©ussie ! Connectez-vous.', 'success')
            return redirect('/login')
        except sqlite3.IntegrityError:
            flash('‚ùå Cet email existe d√©j√†', 'error')
        finally:
            conn.close()
    return render_template('register.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('‚úÖ D√©connexion r√©ussie', 'success')
    return redirect('/login')

# Route pour le technicien
@app.route('/technicien', methods=['GET', 'POST'])
def technicien():
    if session.get('role') != 'technicien':
        flash('‚ùå Acc√®s refus√©', 'error')
        return redirect('/login')
    
    conn = sqlite3.connect('database.db')
    c = conn.cursor()
    c.execute("SELECT id, full_name FROM users WHERE role = 'medecin'")
    medecins = c.fetchall()
    conn.close()
    
    if request.method == 'POST':
        try:
            nom = request.form['nom']
            prenom = request.form['prenom']
            age = request.form['age']
            diagnostic = request.form['diagnostic']
            examen_demande = request.form['examen_demande']
            medecin_id = request.form['medecin_id'] or None
            
            patient_id_str = f"P{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}"
            
            conn = sqlite3.connect('database.db')
            c = conn.cursor()
            c.execute("""INSERT INTO patients (nom, prenom, age, diagnostic, examen_demande, id_patient, date_naissance, telephone) 
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                      (nom, prenom, age, diagnostic, examen_demande, patient_id_str, "1900-01-01", ""))
            patient_id = c.lastrowid
            
            files = request.files.getlist('dicom_files')
            uploaded_paths = []
            patient_folder = os.path.join('uploads', str(patient_id))
            os.makedirs(patient_folder, exist_ok=True)

            if not files:
                flash('‚ùå Aucun fichier DICOM s√©lectionn√©', 'error')
                return redirect('/technicien')

            for file in files:
                if file and file.filename:
                    clean_filename = file.filename.replace(' ', '_').replace('/', '_').replace('\\', '_')
                    filepath = os.path.join(patient_folder, clean_filename)
                    
                    counter = 1
                    original_filepath = filepath
                    while os.path.exists(filepath):
                        name, ext = os.path.splitext(original_filepath)
                        filepath = f"{name}_{counter}{ext}"
                        counter += 1

                    file.save(filepath)
                    uploaded_paths.append(filepath)

            if not uploaded_paths:
                flash('‚ùå Aucun fichier DICOM valide trouv√©', 'error')
                return redirect('/technicien')

            dicom_path_str = ','.join(uploaded_paths)
            uploaded_at = datetime.now().isoformat()
            
            c.execute("""INSERT INTO exams (patient_id, technicien_id, medecin_id, status, dicom_path, uploaded_at) 
                        VALUES (?, ?, ?, ?, ?, ?)""",
                      (patient_id, session['user_id'], medecin_id, 'Attribu√©', dicom_path_str, uploaded_at))
            
            exam_id = c.lastrowid
            conn.commit()
            conn.close()
            
            flash(f'‚úÖ Patient {patient_id_str} cr√©√© et examen #{exam_id} upload√© avec succ√®s !', 'success')
            return redirect('/technicien')
        
        except Exception as e:
            flash(f'‚ùå Erreur : {str(e)}', 'error')
    
    return render_template('technicien.html', medecins=medecins)

# ‚úÖ Route pour servir les fichiers DICOM
@app.route('/download/<path:filepath>')
def download_file(filepath):
    decoded_path = urllib.parse.unquote(filepath)
    if '..' in decoded_path:
        return "Acc√®s refus√©", 403
    full_path = os.path.join(app.root_path, decoded_path)
    if not os.path.exists(full_path):
        return "Fichier non trouv√©", 404
    return send_from_directory(os.path.dirname(full_path), os.path.basename(full_path), mimetype='application/dicom')

# ‚úÖ Route pour visualiser ‚Äî G√©n√®re un .bat avec encodage Windows
@app.route('/visualiser/<int:exam_id>')
def visualiser(exam_id):
    if session.get('role') != 'medecin':
        flash('‚ùå Acc√®s refus√©', 'error')
        return redirect('/login')
    
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("""
        SELECT e.*, p.nom as patient_nom, p.prenom as patient_prenom, p.id_patient as patient_id_patient
        FROM exams e
        JOIN patients p ON e.patient_id = p.id
        WHERE e.id = ? AND e.medecin_id = ?
    """, (exam_id, session['user_id']))
    exam = c.fetchone()
    
    if not exam:
        flash('‚ùå Examen non trouv√© ou non attribu√©', 'error')
        return redirect('/medecin')
    
    # Mettre √† jour visualized_at
    c.execute("UPDATE exams SET visualized_at = ? WHERE id = ?", (datetime.now().isoformat(), exam_id))
    conn.commit()
    conn.close()
    
    if exam['dicom_path']:
        dicom_paths = exam['dicom_path'].split(',')
        valid_paths = [p for p in dicom_paths if os.path.exists(p)]
        
        if not valid_paths:
            flash('‚ùå Aucun fichier DICOM trouv√©', 'error')
            return redirect('/medecin')
        
        # URL publique de ton app d√©ploy√©e
        base_url = "https://telerradiologie-pdo.onrender.com"
        
        # Cr√©er un fichier .bat avec encodage Windows
        bat_content = f"""@echo off
chcp 1252 > nul
echo.
echo ========================================
echo üöÄ D√©marrage du script de visualisation
echo ========================================
echo.

echo ‚úÖ Cr√©ation du dossier temporaire...
mkdir "C:\\temp\\exam_{exam_id}" 2>nul

echo.
echo ========================================
echo üîΩ T√©l√©chargement des fichiers DICOM
echo ========================================
echo.

"""
        for i, path in enumerate(valid_paths):
            filename = os.path.basename(path)
            download_url = f"{base_url}/download/{urllib.parse.quote(path.replace('\\\\', '/'))}"
            bat_content += f'echo T√©l√©chargement de {filename}...\n'
            bat_content += f'powershell -Command "Invoke-WebRequest \'{download_url}\' -OutFile \'C:\\temp\\exam_{exam_id}\\{filename}\'"\n'
            bat_content += f'if errorlevel 1 (\n'
            bat_content += f'    echo ‚ùå √âchec du t√©l√©chargement de {filename}\n'
            bat_content += f'    pause\n'
            bat_content += f'    exit /b 1\n'
            bat_content += f') else (\n'
            bat_content += f'    echo ‚úÖ {filename} t√©l√©charg√©\n'
            bat_content += f')\n'
            bat_content += f'echo.\n'
        
        bat_content += f'''
echo.
echo ========================================
echo üîç Recherche de RadiAnt
echo ========================================
echo.

set "radiant_path="

if exist "C:\\Program Files\\RadiAntViewer64bit\\RadiAntViewer.exe" (
    set "radiant_path=C:\\Program Files\\RadiAntViewer64bit\\RadiAntViewer.exe"
    echo ‚úÖ RadiAnt trouv√© dans Program Files
)

if exist "C:\\Program Files (x86)\\RadiAnt DICOM Viewer\\RadiAntViewer.exe" (
    set "radiant_path=C:\\Program Files (x86)\\RadiAnt DICOM Viewer\\RadiAntViewer.exe"
    echo ‚úÖ RadiAnt trouv√© dans Program Files (x86)
)

if exist "C:\\RadiAnt\\RadiAntViewer.exe" (
    set "radiant_path=C:\\RadiAnt\\RadiAntViewer.exe"
    echo ‚úÖ RadiAnt trouv√© dans C:\\RadiAnt
)

if defined radiant_path (
    echo.
    echo ========================================
    echo üéâ Lancement de RadiAnt
    echo ========================================
    echo.
    echo Ouvrir : %radiant_path%
    "%radiant_path%" "C:\\temp\\exam_{exam_id}"
    if errorlevel 1 (
        echo ‚ùå √âchec du lancement de RadiAnt
        pause
    ) else (
        echo ‚úÖ RadiAnt lanc√© avec succ√®s
    )
) else (
    echo.
    echo ========================================
    echo ‚ùå ERREUR : RadiAnt introuvable
    echo ========================================
    echo.
    echo Veuillez installer RadiAnt depuis :
    echo https://www.radiantviewer.com/
    echo.
    pause
)
'''
        
        # ‚úÖ Forcer l'encodage Windows
        response = make_response(bat_content.encode('cp1252'))
        response.headers['Content-Type'] = 'text/plain; charset=cp1252'
        response.headers['Content-Disposition'] = f'attachment; filename=visualiser_exam_{exam_id}.bat'
        return response
    
    flash('‚ùå Aucun fichier DICOM associ√©', 'error')
    return redirect('/medecin')

# Route pour r√©diger le rapport
@app.route('/rapport/<int:exam_id>', methods=['GET', 'POST'])
def rediger_rapport(exam_id):
    if session.get('role') != 'medecin':
        flash('‚ùå Acc√®s refus√©', 'error')
        return redirect('/login')
    
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("""
        SELECT e.*, p.nom as patient_nom, p.prenom as patient_prenom, p.id_patient as patient_id_patient
        FROM exams e
        JOIN patients p ON e.patient_id = p.id
        WHERE e.id = ? AND e.medecin_id = ?
    """, (exam_id, session['user_id']))
    exam = c.fetchone()
    
    if not exam:
        flash('‚ùå Examen non trouv√© ou non attribu√©', 'error')
        return redirect('/medecin')
    
    if request.method == 'POST':
        try:
            date_examen = request.form['date_examen']
            age = request.form['age']
            examen_demande = request.form['examen_demande']
            technique = request.form['technique']
            resultat = request.form['resultat']
            conclusion = request.form['conclusion']
            
            # G√©n√®re le document Word
            doc = Document()
            section = doc.sections[0]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            
            header = doc.sections[0].header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = f"Gagnoa le {date_examen}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # R√©cup√©rer les infos du patient
            c.execute("SELECT * FROM patients WHERE id = ?", (exam['patient_id'],))
            patient = c.fetchone()
            
            doc.add_paragraph(f"Nom : {patient['nom'].upper()}      Pr√©noms : {patient['prenom'].upper()}")
            doc.add_paragraph(f"Age : {patient['age']} ANS")
            doc.add_paragraph(f"Diagnostic : {patient['diagnostic']}")
            doc.add_paragraph(f"Examen demand√© : {patient['examen_demande']}")
            doc.add_paragraph()
            
            tech_para = doc.add_paragraph()
            tech_para.add_run("TECHNIQUE : ").bold = True
            tech_para.add_run(technique)
            
            doc.add_paragraph()
            
            result_para = doc.add_paragraph()
            result_para.add_run("¬∑  RESULTAT\n").bold = True
            result_para.add_run(resultat)
            
            doc.add_paragraph()
            
            concl_para = doc.add_paragraph()
            concl_para.add_run("¬∑  CONCLUSION\n").bold = True
            concl_para.add_run(conclusion)
            
            doc.add_paragraph()
            
            sign_para = doc.add_paragraph()
            sign_para.add_run(f"DR {session['full_name']}\n").bold = True
            sign_para.add_run("Radiologue").italic = True
            
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = "Ce document est la propri√©t√© de la clinique de Gagnoa - Reproduction interdite"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            rapport_folder = os.path.join('rapports', str(exam['patient_id']))
            os.makedirs(rapport_folder, exist_ok=True)
            word_filename = f"rapport_{exam_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            word_filepath = os.path.join(rapport_folder, word_filename)
            doc.save(word_filepath)
            
            # Met √† jour la base de donn√©es
            c.execute("""
                UPDATE exams 
                SET rapport_path = ?, status = 'Termin√©', rapport_submitted_at = ?
                WHERE id = ?
            """, (word_filepath, datetime.now().isoformat(), exam_id))
            conn.commit()
            conn.close()
            
            flash('‚úÖ Rapport g√©n√©r√© et envoy√© √† la clinique !', 'success')
            return redirect('/medecin')
            
        except Exception as e:
            flash(f'‚ùå Erreur lors de la g√©n√©ration du rapport : {str(e)}', 'error')
    
    default_values = {
        'date_examen': datetime.now().strftime('%d %B %Y').upper(),
        'age': '54',
        'examen_demande': 'SCANNER DES MEMBRES INFERIEURS',
        'technique': 'Acquisition volumique sans et avec injection de contraste sur les membres inf√©rieurs et reconstruction multiplanaires.',
        'resultat': 'Mise en √©vidence d\'une ost√©on√©crose de la t√™te f√©morale gauche qui pr√©sente des zones de lacunes.\nEn regard de la t√™te, on note des clart√©s a√©riques de type emphys√©mateux.\nLymph≈ìd√®me du membre inf√©rieur gauche sans collection abc√©d√©e.\nPerm√©abilit√© des structures vasculaires',
        'conclusion': 'Osteo-Arthrite De La Hanche Gauche Probablement En Rapport Avec Des Germes Ana√©robie Avec Lymph≈ìd√®me De La Jambe.'
    }
    
    return render_template('rapport.html', exam=exam, defaults=default_values)

# Route pour le m√©decin
@app.route('/medecin')
def medecin():
    if session.get('role') != 'medecin':
        flash('‚ùå Acc√®s refus√©', 'error')
        return redirect('/login')
    
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("""
        SELECT e.*, p.nom as patient_nom, p.prenom as patient_prenom, p.id_patient as patient_id_patient
        FROM exams e
        JOIN patients p ON e.patient_id = p.id
        WHERE e.medecin_id = ?
        ORDER BY e.uploaded_at DESC
    """, (session['user_id'],))
    exams = c.fetchall()
    conn.close()
    
    return render_template('medecin.html', exams=exams)

# Route pour l'agent tireur
@app.route('/agent_tireur')
def agent_tireur():
    if session.get('role') != 'agent_tireur':
        flash('‚ùå Acc√®s refus√©', 'error')
        return redirect('/login')
    
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("""
        SELECT e.*, p.nom as patient_nom, p.prenom as patient_prenom, p.id_patient as patient_id_patient,
               u.full_name as medecin_full_name
        FROM exams e
        JOIN patients p ON e.patient_id = p.id
        JOIN users u ON e.medecin_id = u.id
        WHERE e.status = 'Termin√©' AND e.rapport_printed_at IS NULL
        ORDER BY e.rapport_submitted_at DESC
    """)
    rapports = c.fetchall()
    conn.close()
    
    return render_template('agent_tireur.html', rapports=rapports)

# Route pour la clinique (Admin)
@app.route('/clinique', methods=['GET', 'POST'])
def clinique():
    if session.get('role') != 'admin':
        flash('‚ùå Acc√®s refus√©', 'error')
        return redirect('/login')
    
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    
    # G√©rer la cr√©ation de compte
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'create_user':
            email = request.form['email']
            password = generate_password_hash(request.form['password'])
            full_name = request.form['full_name']
            role = request.form['role']
            try:
                c.execute("INSERT INTO users (email, password, full_name, role) VALUES (?, ?, ?, ?)",
                          (email, password, full_name, role))
                conn.commit()
                flash('‚úÖ Compte cr√©√© avec succ√®s !', 'success')
            except sqlite3.IntegrityError:
                flash('‚ùå Cet email existe d√©j√†', 'error')
        
        elif action == 'delete_user':
            user_id = request.form['user_id']
            c.execute("DELETE FROM users WHERE id = ? AND role != 'admin'", (user_id,))
            conn.commit()
            flash('‚úÖ Compte supprim√© !', 'success')
        
        elif action == 'delete_exam':
            exam_id = request.form['exam_id']
            # Supprimer les fichiers DICOM
            c.execute("SELECT dicom_path FROM exams WHERE id = ?", (exam_id,))
            exam = c.fetchone()
            if exam and exam['dicom_path']:
                for path in exam['dicom_path'].split(','):
                    if os.path.exists(path):
                        os.remove(path)
                # Supprimer le dossier parent s'il est vide
                if exam['dicom_path']:
                    folder = os.path.dirname(exam['dicom_path'].split(',')[0])
                    if os.path.exists(folder) and not os.listdir(folder):
                        os.rmdir(folder)
            # Supprimer l'examen de la base
            c.execute("DELETE FROM exams WHERE id = ?", (exam_id,))
            conn.commit()
            flash('‚úÖ Examen supprim√© !', 'success')
    
    # R√©cup√©rer tous les utilisateurs
    c.execute("SELECT * FROM users ORDER BY role, full_name")
    users = c.fetchall()
    
    # R√©cup√©rer les examens
    c.execute("""
        SELECT 
            e.*, 
            p.nom as patient_nom, 
            p.prenom as patient_prenom, 
            p.id_patient as patient_id_patient,
            t.full_name as technicien_full_name,
            (
                SELECT u.full_name 
                FROM users u 
                WHERE u.id = e.medecin_id AND u.role = 'medecin'
            ) as medecin_full_name,
            (
                SELECT u.full_name 
                FROM users u 
                WHERE u.id = e.agent_tireur_id AND u.role = 'agent_tireur'
            ) as agent_tireur_full_name
        FROM exams e
        JOIN patients p ON e.patient_id = p.id
        LEFT JOIN users t ON e.technicien_id = t.id
        ORDER BY e.uploaded_at DESC
    """)
    exams = c.fetchall()
    conn.close()
    
    return render_template('clinique.html', exams=exams, users=users)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    app.run(host='0.0.0.0', port=port, debug=False)